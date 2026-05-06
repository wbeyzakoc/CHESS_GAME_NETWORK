from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Satranç_Ağ_Programlama_Proje_Raporu.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="B7C2D0"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "E8EEF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, "9BB5D6")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F8FE")
    set_cell_text(cell, f"{title}\n{body}", bold=False)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "244061"),
        ("Heading 3", 11, "000000"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Satranç Ağ Programlama Projesi"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.name = "Arial"
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Java Swing, Socket Programlama ve AWS Tabanlı İki Oyunculu Satranç")
    fp.runs[0].font.name = "Arial"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Satranç Ağ Programlama Projesi")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Java Swing arayüzlü, AWS üzerinde çalışan sunucu ile iki istemcili satranç oyunu")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.italic = True

    add_table(
        doc,
        ["Alan", "Bilgi"],
        [
            ["Ders", "Computer Network Concepts / Ağ Programlama"],
            ["Proje konusu", "Chess"],
            ["Programlama dili", "Java"],
            ["Arayüz teknolojisi", "Java Swing / NetBeans GUI Builder"],
            ["Ağ yapısı", "TCP Socket, istemci-sunucu mimarisi"],
            ["Sunucu ortamı", "AWS EC2, console tabanlı Server.java"],
            ["Varsayılan port", "5003"],
        ],
        [2.1, 4.6],
    )

    doc.add_heading("1. Proje Özeti", level=1)
    doc.add_paragraph(
        "Bu proje, iki oyuncunun ağ üzerinden satranç oynayabildiği Java tabanlı bir oyun uygulamasıdır. "
        "Oyun arayüzü Java Swing ile hazırlanmış, gerçek oyun davranışı game.java ve GameLogic.java sınıflarında "
        "uygulanmıştır. Server.java sınıfı grafik arayüze ihtiyaç duymadan konsol üzerinde çalışır ve iki istemciyi "
        "eşleştirerek hamle mesajlarını oyuncular arasında iletir."
    )
    doc.add_paragraph(
        "Proje, raporda istenen Java kullanımı, başlangıç ekranı, bitiş ekranı, tekrar oynama, iki istemci desteği, "
        "AWS IP adresiyle sunucuya bağlanma ve platform bağımsız kaynak yükleme gereksinimlerini karşılayacak şekilde "
        "düzenlenmiştir."
    )

    doc.add_heading("2. Gereksinimlerin Karşılanması", level=1)
    add_table(
        doc,
        ["Rapor Gereksinimi", "Projede Karşılığı", "Durum"],
        [
            ["Programlama dili Java olmalıdır.", "Tüm istemci, oyun ve sunucu kodları Java ile yazılmıştır.", "Karşılandı"],
            ["Arayüz kullanıcı dostu ve tamamlanmış olmalıdır.", "enter.java, ready.java, game.java ve EndScreen.java ile akış tamamlanmıştır.", "Karşılandı"],
            ["Sunucu grafik arayüzsüz olabilir.", "Server.java konsol uygulamasıdır.", "Karşılandı"],
            ["Sunucu AWS üzerinde çalışmalı ve client AWS IP ile bağlanmalıdır.", "Server.java EC2 üzerinde çalıştırılır; ready.java IP/port girişi alır.", "Karşılandı"],
            ["Uygulama değerlendirici bilgisayarında çalışmalıdır.", "Görseller AssetLoader ile resource klasöründen yüklenir; sabit kişisel path bağımlılığı azaltılmıştır.", "Karşılandı"],
            ["Oyun kapatılmadan tekrar oynanabilmelidir.", "EndScreen ekranındaki Play Again butonu oyunu yeniden başlatır.", "Karşılandı"],
            ["Server birden fazla client desteklemelidir.", "Server iki client kabul eder, renk atar ve hamleleri karşı tarafa relay eder.", "Karşılandı"],
            ["Kod temel programlama ilkelerine uygun olmalıdır.", "Oyun mantığı, bağlantı, varlık modeli ve asset yükleme ayrı sınıflara ayrılmıştır.", "Karşılandı"],
        ],
        [2.4, 3.6, 1.0],
    )

    doc.add_heading("3. Sistem Mimarisi", level=1)
    doc.add_paragraph(
        "Proje istemci-sunucu mimarisiyle çalışır. Her oyuncu kendi bilgisayarında client uygulamasını açar. "
        "Client uygulaması AWS EC2 üzerinde çalışan Server.java uygulamasına TCP socket ile bağlanır. Server ilk iki "
        "bağlantıyı bir maç olarak eşleştirir; birinci oyuncuya beyaz/Türk, ikinci oyuncuya siyah/Yunan rengi atanır."
    )
    add_note(
        doc,
        "Mimari Akış",
        "Client 1 -> AWS Server -> Client 2. Hamleler MOVE,fromX,fromY,toX,toY biçiminde gönderilir. "
        "Server oyun kuralı hesaplamaz; güvenilir iletişim köprüsü olarak iki client arasında mesaj iletir.",
    )
    add_table(
        doc,
        ["Bileşen", "Sorumluluk"],
        [
            ["Client GUI", "Kullanıcıdan taş seçimi ve hedef kare seçimi alır, oyun durumunu gösterir."],
            ["GameLogic", "Satranç hamle doğrulaması, şah, mat, pat, rok ve piyon terfisini yönetir."],
            ["ClientConnection", "Socket bağlantısı üzerinden sunucuya mesaj gönderir ve sunucudan mesaj alır."],
            ["Server", "İki oyuncuyu eşleştirir, renk bilgisi gönderir ve hamleleri diğer client'a aktarır."],
            ["AWS EC2", "Server uygulamasının ortak erişilebilir IP adresiyle çalıştığı bulut ortamıdır."],
        ],
        [1.8, 5.0],
    )

    doc.add_heading("4. Sınıf ve Dosya Açıklamaları", level=1)
    add_table(
        doc,
        ["Dosya", "Açıklama"],
        [
            ["Chess.java", "Programın ana başlangıç noktasıdır; enter ekranını açar."],
            ["enter.java", "İlk tanıtım/karşılama ekranıdır."],
            ["ready.java", "Oyun tanıtımı ve sunucu IP/port giriş ekranıdır."],
            ["game.java", "Satranç tahtası, taş seçimleri, sıra kontrolü, skor/gösterge alanı ve ağdan gelen hamlelerin uygulanmasını içerir."],
            ["GameLogic.java", "Taş hareket kuralları, şah tehdidi, yasal hamle kontrolü, rok, piyon terfisi, mat ve pat analizini içerir."],
            ["Piece.java", "Taş tipi, rengi, koordinatı ve hareket edip etmediği bilgisini tutan model sınıfıdır."],
            ["ClientConnection.java", "TCP Socket, BufferedReader ve PrintWriter kullanarak client-server iletişimini sağlar."],
            ["Server.java", "AWS üzerinde çalışan konsol sunucusudur; iki client'ı eşleştirir ve mesajları aktarır."],
            ["AssetLoader.java", "Görselleri classpath resource klasöründen, gerekirse img klasöründen yükler."],
            ["EndScreen.java", "Oyun sonucu, Play Again ve Exit seçeneklerini gösteren bitiş ekranıdır."],
        ],
        [1.8, 5.0],
    )

    doc.add_heading("5. Oyun Kuralları ve Mantık", level=1)
    add_bullets(
        doc,
        [
            "Oyunda sıra kontrolü uygulanır; beyaz/Türk başlar, ardından sıra siyah/Yunan oyuncuya geçer.",
            "Piyon, kale, at, fil, vezir ve şah için temel satranç hareketleri kontrol edilir.",
            "Kendi taşını yeme ve tahta dışına hamle yapma engellenir.",
            "Bir oyuncunun kendi şahını tehdit altında bırakacak hamlesi geçersiz sayılır.",
            "Şah tehdidi oluştuğunda kullanıcı bilgilendirilir.",
            "Yasal hamle kalmadığında şah tehdidi varsa şah mat, şah tehdidi yoksa pat sonucu üretilir.",
            "Rok hamlesi, şah ve kalenin daha önce hareket etmemesi ve geçiş karelerinin tehdit altında olmaması şartlarıyla desteklenir.",
            "Son sıraya ulaşan piyon otomatik olarak vezire terfi eder ve görseli güncellenir.",
        ],
    )

    doc.add_heading("6. Ağ Haberleşmesi", level=1)
    doc.add_paragraph(
        "İstemci ve sunucu arasındaki haberleşme TCP socket üzerinden yapılır. ClientConnection sınıfı Socket nesnesi "
        "oluşturur; PrintWriter ile mesaj gönderir ve BufferedReader ile mesaj alır. Server.java iki client kabul eder "
        "ve iki yönlü relay thread'leriyle mesajları karşı tarafa aktarır."
    )
    add_table(
        doc,
        ["Mesaj", "Anlamı"],
        [
            ["COLOR,beyaz", "Sunucu tarafından birinci client'a beyaz/Türk rolü atanır."],
            ["COLOR,siyah", "Sunucu tarafından ikinci client'a siyah/Yunan rolü atanır."],
            ["INFO,Rakip bekleniyor...", "İlk client bağlandığında ikinci oyuncunun beklenmesi gerektiğini bildirir."],
            ["MOVE,x1,y1,x2,y2", "Bir taşın kaynak koordinattan hedef koordinata hamle yaptığını bildirir."],
        ],
        [1.8, 5.0],
    )

    doc.add_heading("7. AWS Kurulumu ve Çalıştırma", level=1)
    doc.add_paragraph(
        "Raporda belirtilen zorunluluk nedeniyle sunucu uygulaması AWS EC2 üzerinde çalıştırılmalıdır. "
        "Client bilgisayarlarında Server.java çalıştırılmaz; her iki client da AWS Public IPv4 adresi ve 5003 portu ile bağlanır."
    )
    add_numbered(
        doc,
        [
            "AWS EC2 üzerinde Amazon Linux 2023 tabanlı t3.micro veya t2.micro bir instance oluşturulur.",
            "Security Group inbound kurallarına SSH 22 ve Custom TCP 5003 eklenir.",
            "EC2 Public IPv4 adresi not edilir. Test ortamında kullanılan örnek IP: 51.20.82.77.",
            "Mac terminalinden SSH ile bağlantı kurulur: ssh -i ~/.ssh/chess-key.pem ec2-user@<AWS_PUBLIC_IP>.",
            "Java kurulur: sudo dnf install java-17-amazon-corretto java-17-amazon-corretto-devel -y.",
            "Server.java dosyası AWS'ye kopyalanır ve javac ile derlenir.",
            "Sunucu kalıcı çalışsın diye nohup java com.mycompany.chess.client.Server > server.log 2>&1 & komutu kullanılabilir.",
            "Client açılırken Server IP address alanına AWS Public IPv4, port alanına 5003 yazılır.",
        ],
    )
    add_note(
        doc,
        "Bağlantı Kontrolü",
        "Client bağlanamıyorsa önce AWS üzerinde ss -ltnp | grep 5003 komutu ile server'ın dinlediği kontrol edilir. "
        "Yerel bilgisayardan nc -vz <AWS_PUBLIC_IP> 5003 komutu bağlantıyı test etmek için kullanılabilir.",
    )

    doc.add_heading("8. Kullanım Senaryosu", level=1)
    add_numbered(
        doc,
        [
            "AWS üzerinde Server.java çalışır durumda bırakılır.",
            "Birinci oyuncu Chess.java ile client uygulamasını açar.",
            "Start ekranında AWS IP adresi ve 5003 portu girilir.",
            "İkinci oyuncu farklı bilgisayardan aynı AWS IP ve port ile bağlanır.",
            "Server iki oyuncuyu eşleştirir ve terminalde New chess match started mesajı görülür.",
            "Beyaz/Türk ilk hamleyi yapar; hamle ikinci client ekranında görünür.",
            "Sıra siyah/Yunan oyuncuya geçer. Oyun mat, pat veya kullanıcı çıkışıyla tamamlanır.",
            "Bitiş ekranından Play Again ile uygulama kapatılmadan yeni oyun başlatılabilir.",
        ],
    )

    doc.add_heading("9. Test Senaryoları", level=1)
    add_table(
        doc,
        ["Test", "Beklenen Sonuç", "Durum"],
        [
            ["Server başlatma", "AWS terminalinde Chess server started on port 5003 çıktısı görülür.", "Başarılı"],
            ["Port erişimi", "nc -vz <AWS_PUBLIC_IP> 5003 succeeded çıktısı verir.", "Başarılı"],
            ["İki client bağlantısı", "Server terminalinde New chess match started mesajı oluşur.", "Başarılı"],
            ["Sıra kontrolü", "Sırası olmayan oyuncu hamle yapamaz ve uyarı alır.", "Başarılı"],
            ["Hamle aktarımı", "Bir client'ta yapılan hamle diğer client ekranında uygulanır.", "Başarılı"],
            ["Geçersiz hamle", "Kurala uymayan hamle reddedilir.", "Başarılı"],
            ["Piyon hareketi", "Piyon ilk hamlede iki kare, normalde bir kare ilerleyebilir.", "Başarılı"],
            ["Şah kontrolü", "Şahı açıkta bırakan hamle engellenir.", "Başarılı"],
            ["Oyun sonu", "Mat/pat durumunda EndScreen açılır.", "Başarılı"],
        ],
        [1.7, 4.0, 1.0],
    )

    doc.add_heading("10. Güvenlik ve Sınırlamalar", level=1)
    add_bullets(
        doc,
        [
            "AWS Security Group üzerinde 5003 portu internetten erişime açılmıştır; bu proje/demo amacıyla yeterlidir.",
            "Server mesajları relay eder; oyuncu hamle doğrulaması client tarafında yapılır.",
            "Client bağlantısı kesilirse oyuncular oyunu yeniden açıp aynı AWS IP/port ile bağlanmalıdır.",
            "EC2 instance durdurulursa server kapanır ve Public IPv4 adresi değişebilir.",
            "Daha ileri geliştirme için kullanıcı adı, oda sistemi, bağlantı kopunca yeniden bağlanma ve server tarafı hamle doğrulaması eklenebilir.",
        ],
    )

    doc.add_heading("11. Sonuç", level=1)
    doc.add_paragraph(
        "Proje, Java Swing arayüzü ve TCP socket haberleşmesini birleştirerek iki oyunculu bir satranç oyunu sunmaktadır. "
        "AWS EC2 üzerinde çalışan console server sayesinde oyuncular aynı ağda olmak zorunda kalmadan ortak bir sunucu "
        "üzerinden bağlanabilmektedir. Oyun; sıra kontrolü, geçerli hamle doğrulama, şah/mat/pat kontrolü, rok ve piyon "
        "terfisi gibi temel satranç özelliklerini içerdiği için kullanıcı açısından tamamlanmış bir oyun deneyimi sağlar."
    )

    doc.add_page_break()
    doc.add_heading("Ek A - Çalıştırma Komutları", level=1)
    add_table(
        doc,
        ["Amaç", "Komut / Değer"],
        [
            ["AWS'ye bağlanma", "ssh -i ~/.ssh/chess-key.pem ec2-user@51.20.82.77"],
            ["Java kurulumu", "sudo dnf install java-17-amazon-corretto java-17-amazon-corretto-devel -y"],
            ["Server derleme", "javac com/mycompany/chess/client/Server.java"],
            ["Server çalıştırma", "java com.mycompany.chess.client.Server"],
            ["Arka planda çalıştırma", "nohup java com.mycompany.chess.client.Server > server.log 2>&1 &"],
            ["Server kontrol", "ss -ltnp | grep 5003"],
            ["Port testi", "nc -vz 51.20.82.77 5003"],
            ["Client IP", "51.20.82.77"],
            ["Client port", "5003"],
        ],
        [2.0, 4.8],
    )

    doc.add_heading("Ek B - Proje Dosya Yapısı", level=1)
    add_bullets(
        doc,
        [
            "src/main/java/com/mycompany/chess/client: Java kaynak kodları",
            "src/main/resources/com/mycompany/chess/assets: oyun görselleri",
            "pom.xml: Maven proje yapılandırması ve ana sınıf ayarı",
            "img: geliştirme sırasında kullanılan görsel dosyaların yedek klasörü",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
